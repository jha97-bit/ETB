#!/usr/bin/env python3
"""Export ETB-Documentation.md + PROJECT_REPORT.md to .docx and .pdf (project root)."""

from __future__ import annotations

import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
OUT_DOCX = ROOT / "ETB-Documentation-Pack.docx"
OUT_PDF = ROOT / "ETB-Documentation-Pack.pdf"


def _read_parts() -> str:
    parts = []
    for name in ("ETB-Documentation.md", "PROJECT_REPORT.md"):
        p = ROOT / name
        if p.exists():
            parts.append(f"\n\n---\n\n# {name}\n\n")
            parts.append(p.read_text(encoding="utf-8"))
    return "".join(parts) if parts else ""


def _export_docx(text: str, dest: Path) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for raw in text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph()
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.strip() == "---":
            doc.add_paragraph("— — —")
        else:
            doc.add_paragraph(line)

    doc.save(dest)


def _export_pdf(text: str, dest: Path) -> None:
    from io import BytesIO

    import markdown
    from xhtml2pdf import pisa

    body = markdown.markdown(
        text,
        extensions=["tables", "fenced_code", "nl2br"],
    )
    html = f"""<!DOCTYPE html><html><head><meta charset="utf-8"/>
    <style>
    body {{ font-family: Helvetica, Arial, sans-serif; font-size: 11pt; margin: 2cm; }}
    h1 {{ font-size: 18pt; }}
    h2 {{ font-size: 14pt; margin-top: 14pt; }}
    h3 {{ font-size: 12pt; margin-top: 10pt; }}
    code, pre {{ font-size: 9pt; }}
    table {{ border-collapse: collapse; width: 100%; }}
    th, td {{ border: 1px solid #ccc; padding: 4px; text-align: left; }}
    </style></head><body>{body}</body></html>"""

    buf = BytesIO()
    status = pisa.CreatePDF(html.encode("utf-8"), dest=buf, encoding="utf-8")
    if status.err:
        raise RuntimeError(f"xhtml2pdf reported {status.err} errors")
    dest.write_bytes(buf.getvalue())


def main() -> int:
    text = _read_parts()
    if not text.strip():
        print("No source markdown found.", file=sys.stderr)
        return 1

    try:
        _export_docx(text, OUT_DOCX)
        print(f"Wrote {OUT_DOCX}")
    except Exception as e:
        print(f"docx failed: {e}", file=sys.stderr)
        return 1

    try:
        _export_pdf(text, OUT_PDF)
        print(f"Wrote {OUT_PDF}")
    except Exception as e:
        print(f"pdf failed: {e}", file=sys.stderr)
        return 1

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
